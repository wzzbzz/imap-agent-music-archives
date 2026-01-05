import os
import json
import zipfile
import io
import shutil
import tempfile

from docx import Document

from utils import clean_text, sanitize_for_json, get_release_number_fallback, slugify_filename, is_already_downloaded, mark_as_downloaded

from normalize_audio import normalize_audio
from config import IMAP_SERVER,EMAIL_USER,EMAIL_PASS,SENDER_EMAIL,EMAIL_SUBJECT

from imap_utils import fetch_emails

BASE_DIR = ""

def process_zip_attachment(att_payload, target_dir):
    """Unzips files and extracts text from any .docx found."""
    saved_files = []
    lyrics_text = ""
    with zipfile.ZipFile(io.BytesIO(att_payload)) as z:
        for z_info in z.infolist():
            filename = os.path.basename(z_info.filename)
            # Skip system junk
            if not filename or filename.startswith('__') or filename.startswith('.'):
                continue
            
            # Slugify the filename for the audio directory
            slugged_name = slugify_filename(clean_text(filename))
            target_path = os.path.join(target_dir, slugged_name)
            
            with z.open(z_info.filename) as source, open(target_path, 'wb') as target:
                target.write(source.read())
            
            saved_files.append({"original": filename, "slugified": slugged_name})

            # Extract lyrics if it's a docx
            if slugged_name.lower().endswith('.docx'):
                try:
                    doc = Document(target_path)
                    lyrics_text = "\n".join([p.text for p in doc.paragraphs])
                except:
                    pass
            
            # Normalize if it's audio
            if slugged_name.lower().endswith(('.mp3', '.wav', '.m4a')):
                normalize_audio(target_path)
                
    return saved_files, lyrics_text

                
def process_zip_attachment(att_payload, target_dir):
    """
    Unzips files to target_dir. 
    Returns: (list_of_filenames, extracted_lyrics_text)
    """
    saved_files = []
    lyrics_text = ""
    
    # Open the zip in memory
    with zipfile.ZipFile(io.BytesIO(att_payload)) as z:
        for file_info in z.infolist():
            # Filter out system junk (__MACOSX, .DS_Store, etc)
            filename = os.path.basename(file_info.filename)
            if not filename or filename.startswith('.') or filename.startswith('__'):
                continue
                
            # Define the save path
            file_path = os.path.join(target_dir, filename)
            
            # Extract and save the file
            with z.open(file_info.filename) as source, open(file_path, 'wb') as target:
                target.write(source.read())
            
            saved_files.append(filename)
            
            # EDGE CASE: If it's the lyrics file, extract the text now
            if filename.lower().endswith('.docx'):
                try:
                    doc = Document(file_path)
                    lyrics_text = "\n".join([para.text for para in doc.paragraphs])
                    print(f"📝 Extracted lyrics from {filename}")
                except Exception as e:
                    print(f"⚠️ Could not read docx {filename}: {e}")

    return saved_files, lyrics_text

def manually_process_local_zip(zip_path, release_number):

    # 1. Setup the paths
    final_issue_dir = os.path.join(BASE_DIR, f"Issue_{release_number}")
    audio_dir = os.path.join(final_issue_dir, "audio")
    os.makedirs(audio_dir, exist_ok=True)
    
    print(f"📦 Manually processing local ZIP for Issue {release_number}...")
    
    # 2. Read the file into memory as a payload
    with open(zip_path, 'rb') as f:
        zip_payload = f.read()
    
    # 3. Use your existing function
    zip_files, lyrics = process_zip_attachment(zip_payload, audio_dir)
    
    # 4. Create a minimal raw.json for this manual entry
    manual_json = {
        "uid": "MANUAL",
        "subject": f"Manual Import: Issue #{release_number}",
        "body": lyrics,
        "attachments": zip_files,
        "note": "Imported manually because IMAP fetch skipped the ZIP."
    }
    
    with open(os.path.join(final_issue_dir, "raw.json"), "w") as f:
        import json
        json.dump(manual_json, f, indent=4)
        
    print(f"✅ Successfully processed {len(zip_files)} files from local ZIP.")

def process_downloaded_audio(release_number):
    issue_dir = os.path.join(BASE_DIR, f"Issue_{release_number}")
    raw_json_path = os.path.join(issue_dir, "raw.json")
    audio_dir = os.path.join(issue_dir, "audio")

    if not os.path.exists(raw_json_path) or not os.path.exists(audio_dir):
        print(f"❌ Missing required files for Issue {release_number}")
        return

    with open(raw_json_path, 'r') as f:
        data = json.load(f)

    new_attachment_map = []
    
    # Process every audio file found
    for filename in os.listdir(audio_dir):
        if filename.lower().endswith((".mp3", ".m4a", ".wav", ".ogg", ".aac")):
            input_path = os.path.join(audio_dir, filename)
            clean_name = slugify_filename(filename)
            
            print(f"🚀 Processing: {filename} -> {clean_name}")
            
            # Run the normalization and conversion logic
            success = normalize_audio_to_mp3(input_path, clean_name)
            
            if success:
                new_attachment_map.append({
                    "original_name": filename,
                    "clean_name": clean_name,
                    "path": f"audio/{clean_name}",
                    "type": "audio/mpeg"
                })

    # Update raw.json
    data["attachments"] = new_attachment_map
    with open(raw_json_path, 'w') as f:
        json.dump(data, f, indent=4)
    
    print(f"✅ Issue {release_number} fully normalized and mapped.")

def normalize_audio_to_mp3(input_path, output_filename):
    """Normalizes to -16 LUFS and forces MP3 output."""
    output_path = os.path.join(os.path.dirname(input_path), output_filename)
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        temp_output = os.path.join(tmp_dir, f"norm_{output_filename}")
        
        norm = FFmpegNormalize(
            normalization_type='ebu',
            target_level=-16,
            audio_codec='libmp3lame',
            output_format='mp3',
            extra_output_options=['-b:a', '320k'],
            print_stats=False
        )
        
        try:
            media_file = MediaFile(norm, input_path, temp_output)
            norm.media_files.append(media_file)
            norm.run_normalization()
            
            if os.path.exists(temp_output) and os.path.getsize(temp_output) > 0:
                # Remove original if it has a different name/extension to prevent clutter
                if os.path.abspath(input_path) != os.path.abspath(output_path):
                    os.remove(input_path)
                
                shutil.move(temp_output, output_path)
                return True
        except Exception as e:
            print(f"❌ Error: {e}")
            return False
    return False

def check_if_issues_are_present():
    print(f"{'ISSUE':<10} | {'FILE STATUS'}")
    print("-" * 50)

    issue_folders = [d for d in os.listdir(BASE_DIR) if d.startswith("Issue_")]
    issue_folders.sort(key=lambda x: int(x.split('_')[1]) if x.split('_')[1].isdigit() else 0)


    release_numbers = {int(f.split("_")[1]) for f in issue_folders if f.split("_")[1].isdigit()}
    earliest_issue = min(release_numbers)
    last_issue = max(release_numbers)

    valid = True
    missing = []
    for x in range(earliest_issue, last_issue+1):
        
        if not issue_is_present(x):
            valid=False
            missing.append(x)

    if valid==False:
        print(f'missing issues: {str(missing)}')
    else:
        print(f"all issues present")

def issue_directory(release_number):
    return os.path.join(BASE_DIR, f"Issue_{release_number}")

def issue_is_present(release_number):
    return os.path.exists(issue_directory(release_number))

def check_if_attachments_are_present(release_number):
    if not issue_is_present(release_number):
        print(f"issue {release_number} not present")
        return False

    raw_json = os.path.join(issue_directory(release_number), 'raw.json')
    with open(raw_json, 'rb') as episode_json:
        episode_data = json.load(episode_json)

def process_single_email(arguments, force=False):

    msg = arguments.get('msg')
    ri = arguments.get('release_indicator')
    base_dir = arguments.get('base_dir')

    print(f"📥 Fetching UID {msg.uid}: {msg.subject}")
    
    clean_subj = clean_text(msg.subject)
    print(f"Subject: {clean_subj}")
    release_num_guess = get_release_number_fallback(clean_subj)

    final_issue_dir = os.path.join(arguments.get('base_dir'), f"{ri}_{release_num_guess}")

    print(f"process_single_email: final_issue_dir {final_issue_dir} ")


    if not force and os.path.exists(final_issue_dir) :
        print(f"⏭️  {ri} {release_num_guess} already exists. Skipping.")
        return

    if not str(release_num_guess).isdigit():
        print(f'    ❌ Error parsing {clean_subj}: "{release_num_guess}" is not a valid number')
        return

    try:
        
        print(f"\n📄 New {ri} Found: {clean_subj}")
        
        attachment_map = [{"original": clean_text(att.filename), "slugified": slugify_filename(clean_text(att.filename))} for att in msg.attachments]

        safe_body = sanitize_for_json(msg.text or msg.html)
            
        audio_dir = os.path.join(final_issue_dir, "audio")
        os.makedirs(audio_dir, exist_ok=True)
        
        issue_num = release_num_guess
        audio_dir = os.path.join(final_issue_dir, "audio")
        os.makedirs(audio_dir, exist_ok=True)

        print(f"📥 Archiving {ri} {issue_num}...")

        attachment_map = []
        extracted_lyrics = ""

        for att in msg.attachments:

            orig_name = clean_text(att.filename)

            if orig_name.lower().endswith('.zip'):
                print(f"📦 Unzipping contents of: {orig_name}")
                zip_files, lyrics = process_zip_attachment(att.payload, audio_dir)
                attachment_map.extend(zip_files)
                if lyrics:
                    extracted_lyrics += f"\n{lyrics}"
                    safe_body = lyrics
            else:

                # Standard File Logic
                slugged_name = slugify_filename(orig_name)
                file_path = os.path.join(audio_dir, slugged_name)
                
                with open(file_path, 'wb') as f:
                    f.write(att.payload)
                    
                attachment_map.append({"original": orig_name, "slugified": slugged_name})
                
                if slugged_name.lower().endswith(('.mp3', '.wav', '.m4a')):
                    normalize_audio(file_path)

        #######
        # Define the path to your raw metadata
        # Define the path to your raw metadata
        raw_json_path = os.path.join(final_issue_dir, "raw.json")

        # 1. Prepare current email data
        new_data = {
            "uid": msg.uid,
            "message_id": msg.obj.get('Message-ID'),
            "subject": clean_subj,
            "body": safe_body,
            "date": str(msg.date),
            "from": msg.from_,
            "to": msg.to,
            "attachments": attachment_map # Assuming this is a list of dicts
        }

        if arguments.get("merge") and os.path.exists(raw_json_path):
            with open(raw_json_path, "r") as f:
                try:
                    existing_raw = json.load(f)
                except json.JSONDecodeError:
                    existing_raw = {}

            print(f"🔗 Explicitly merging fragment into existing raw.json...")

            # --- THE EXPLICIT MERGE LOGIC ---

            # UID & Message ID & Subject: Convert to / Append to List
            for list_key in ["uid", "message_id", "subject"]:
                current_val = existing_raw.get(list_key)
                new_val = new_data.get(list_key)
                
                if not isinstance(current_val, list):
                    existing_raw[list_key] = [current_val] if current_val is not None else []
                
                # Only append if it's not already in the list (prevents duplicates on re-runs)
                if new_val not in existing_raw[list_key]:
                    existing_raw[list_key].append(new_val)

            # Body: Simple String Append
            # Adds a newline between parts for readability in the UI
            existing_body = existing_raw.get("body", "")
            new_body = new_data.get("body", "")
            if new_body not in existing_body: # Avoid double-appending the same body
                existing_raw["body"] = f"{existing_body}\n\n--- PART 2 ---\n\n{new_body}".strip()

            # Attachments: Append new map into the old array
            # We assume 'attachments' in existing_raw is already a list
            current_attachments = existing_raw.get("attachments", [])
            if isinstance(current_attachments, str):
                # Fallback if it was previously saved as a JSON string
                current_attachments = json.loads(current_attachments)
            
            # Extend the list with new attachments
            current_attachments.extend(new_data.get("attachments", []))
            existing_raw["attachments"] = current_attachments

            final_output = existing_raw

        else:
            # --- CREATE NEW ---
            print(f"🆕 Creating new raw.json...")
            # Initialize list fields even for single entries to keep the API consistent
            new_data["uid"] = [new_data["uid"]]
            new_data["message_id"] = [new_data["message_id"]]
            new_data["subject"] = [new_data["subject"]]
            final_output = new_data

        # Final Write
        with open(raw_json_path, "w") as f:
            json.dump(final_output, f, indent=4)


    except Exception as e:
        print(f"❌ Error on {clean_subj}: {e}")

def fetch_sonic_twist_emails( arguments ):
    
    BASE_DIR = arguments.get('base_dir',"temp")
    os.makedirs(BASE_DIR, exist_ok=True)

    for msg in fetch_emails(arguments):
        registry_path = os.path.join(BASE_DIR,'downloaded_uids.json')
        if is_already_downloaded(msg.uid, registry_path):
            print(f"⏭️  UID {msg.uid}, {msg.subject} already in registry. Skipping.")
            continue
        
        email_arguments = {"msg":msg,"release_indicator":arguments.get('release_indicator'),"base_dir":arguments.get('base_dir','temp')}
        process_single_email(email_arguments, True)

def retry_email(arguments):
    for msg in fetch_emails(arguments):
        email_arguments = {"msg":msg,"release_indicator":arguments.get('release_indicator'),"base_dir":arguments.get('base_dir','temp'),"merge":arguments.get("merge",False)}
        process_single_email(email_arguments, True)

if __name__ == "__main__":
    # Execute for Issue 60
    # (Make sure the path to your downloaded zip is correct)
    #manually_process_local_zip("/Users/jamespwilliams/Projects/python/sonic_twist_archives/Issue_60/sonic-twist-60.zip", "60")
    #target_missing_email('CA+yqSP4yJreCeh1m9_ncf0vhPLmbJGytE5jN7B7bQGCw5HMxgg@mail.gmail.com')
    #repair_and_upgrade_issue("38")
    #list_attachment_status()
    #check_if_attachments_are_present("30")
    arguments = {
        "base_dir":"even_more_cake_archives",
        "release_indicator":"Volume",
        "release_number": 6,
        "message_id":"CAMyHBL2Xx+SO1T9vkwoSyDLT7jktHiXZea1d6P7=UZrS+_srKA@mail.gmail.com",
        "folder":"[Gmail]/All Mail",
        "sender":"alvyhall@aol.com",
        "subject":"Sonic Twist",
        "before":"12/31/25",
        #"after":"12/31/2024",
        "attachments":True,
        "exclude":("re:","fwd:"),
        "merge": True
        
    }
    #retry_email(arguments)
    process_downloaded_audio( arguments )


